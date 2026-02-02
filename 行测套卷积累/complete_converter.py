import os
import re
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

# ==================== 配置区域 ====================

# 1. 触发【开始删除】的关键词
START_KEYWORDS = ['【答案】', '【解析】', '【拓展】', '【来源】', '正确答案', '参考答案']

# 2. 【强制删除】的前缀特征
# 即使脚本判断“已经进入下一题了”，如果遇到这些开头，依然强制删掉
# 这能解决“误判停止”后，残留的尾部解析 (特别是“故...”系列)
FORCE_DELETE_PREFIXES = [
    '因此，选择', '因此选择', '故本题选', '故正确答案',
    '第一步，', '第二步，', '第三步，', '第四步，',
    'A项：', 'B项：', 'C项：', 'D项：', 'A项 ', 'B项 ',
    '①', '②', '③', '④' # 解析里常出现的序号
]

# 新增：强力删除关键词 (只要包含这些词，整行立刻删除)
# 用于解决 "故本题选" 出现在句子中间或有干扰字符的情况
STRONG_DELETE_CONTAIN = [
    '故本题选', '故正确答案', '故本题正确答案'
]

# 3. 题目序号正则 (兼容多种格式)
# 修正 transfer.py 漏掉 "127)" 这种格式的问题
# 增加对纯数字+空格的支持（防止漏题导致后续误删）
Q_PATTERN = re.compile(r'^\s*(\d+)[、\.．\)\s]')

# 4. 大标题/材料保护正则 (遇到这些绝对停止删除)
HEADER_PATTERN = re.compile(
    r'^\s*第[一二三四五六七八九十]+部分|'  # 第一部分
    r'^\s*[一二三四五六七八九十]+、|'      # 一、常识判断
    r'^\s*根据.*(材料|回答|短文)'          # 根据...回答
)

# ================================================

def iter_block_items(parent):
    """递归遍历文档内容"""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Parent object error")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def has_image(paragraph):
    """
    检查段落是否包含图片 (drawing, pict, object)
    """
    drawings = paragraph._element.findall('.//w:drawing', namespaces=paragraph._element.nsmap)
    if drawings: return True
    picts = paragraph._element.findall('.//w:pict', namespaces=paragraph._element.nsmap)
    if picts: return True
    # 有些公式或嵌入对象可能是 object
    objects = paragraph._element.findall('.//w:object', namespaces=paragraph._element.nsmap)
    if objects: return True
    return False

def check_start_trigger(text):
    """检查是否触发删除模式"""
    for kw in START_KEYWORDS:
        if kw in text:
            return kw
    return None

def is_valid_next_question(found_num, last_num):
    """
    【核心逻辑】验证找到的数字是不是合理的下一题序号
    防止解析里的 "1. xxx" 打断删除
    """
    # 1. 如果还没找到过题目 (last_num=0)，那 1 肯定是题目
    # 如果一开始就找到个 100，也可能是题目（比如只有一部分）
    if last_num == 0:
        return True
    
    # 2. 如果找到的数字 比 上一题数字 小 (例如上一题100，这里发现了1)
    # 可能是解析里的序号，或者是新的一部分重新开始编号
    if found_num < last_num:
        # 如果是1，且距离上一题很远，可能是新部分（通常有大标题保护，这里作为兜底）
        # 但如果是 1，且上一题是 5，那可能是解析里的 1. 2. 3.
        if found_num == 1 and last_num > 10:
            return True # 允许重置为 1
        return False # 其他倒序视为干扰
        
    # 3. 如果跳跃太大 (例如上一题5，这里发现了100)
    # 允许漏题（比如删了几道题），但允许跳跃
    if found_num - last_num > 20:
        return False
        
    return True

def clean_docx_complete(input_path, output_path):
    print(f"处理文件: {os.path.basename(input_path)}")
    doc = Document(input_path)
    
    # 0. 去除标题中的 "解析"
    if len(doc.paragraphs) > 0:
        title_p = doc.paragraphs[0]
        if "解析" in title_p.text:
            title_p.text = title_p.text.replace("解析", "")

    # 状态变量
    is_deleting = False
    last_q_num = 0  # 记录上一个确认为题目的序号
    
    blocks = list(iter_block_items(doc))
    
    # === 第一轮：逻辑擦除 ===
    for block in blocks:
        # 获取段落列表 (无论是直接段落还是表格内的)
        paragraphs = []
        if isinstance(block, Paragraph):
            paragraphs.append(block)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        
        for p in paragraphs:
            text = p.text.strip()
            
            # 【重要修复】如果段落为空 (text=="")，以前是 continue 跳过
            # 这导致如果我们在“删除模式”下，遇到只有图片的段落，它被跳过处理 -> 图片保留 -> 错误！
            # 现在：如果 text 为空，也得检查是不是在删除模式。
            
            if not text:
                if is_deleting:
                    # 如果由于解析里的图片导致段落看起来是空的，但我们处于删除模式，必须删！
                    # p.text = "" 对图片段落可能无效（因为本来就没字），但如果是 mixed，会清空字。
                    # 为了删除图片，最彻底的是清空 runs 或 xml，但这里我们标记它。
                    # 实际上 python-docx 的 p.text="" 会移除所有 runs，包括图片 run。
                    # 所以如果 is_deleting，即使是空文本段落，也要再置空一次？
                    # 不，p.text="" 对空段落无操作。
                    # 我们需要更强力的删除：如果 is_deleting，且是空文本（可能有图），我们得标记这一段需要被物理移除，
                    # 或者简单点：如果确定在 Answer 区，直接 p.clear() ? (docx没有clear)
                    # p.text = "" 基本能清空。
                    # 关键是：如果 continue 了，后面的 "if is_deleting: p.text=''" 就没执行。
                    # 所以删除 continue，往下走。
                    pass 
                else: 
                    # 正常模式下，空行留给第二轮物理清洗处理
                    continue

            # -------------------------------------------------------
            # 1. 优先检查：大标题/材料 (最高优先级停止删除)
            # -------------------------------------------------------
            if HEADER_PATTERN.match(text):
                is_deleting = False
                continue

            # -------------------------------------------------------
            # 2. 检查：题目序号 -> 停止删除
            # -------------------------------------------------------
            match = Q_PATTERN.match(text)
            if match:
                try:
                    found_num = int(match.group(1))
                    if is_valid_next_question(found_num, last_q_num):
                        is_deleting = False
                        last_q_num = found_num
                    else:
                        # 虽像题目但逻辑不对，视为解析内容，继续删除
                        pass
                except ValueError:
                    pass

            # -------------------------------------------------------
            # 3. 检查：开始删除触发器 (答案/解析)
            # -------------------------------------------------------
            keyword = check_start_trigger(text)
            if keyword:
                is_deleting = True
                if text.startswith(keyword):
                    p.text = "" # 彻底清空，防止残留图片
                else:
                    parts = p.text.split(keyword)
                    p.text = parts[0].strip()
                continue

            # -------------------------------------------------------
            # 4. 检查：强制删除特征 (补刀 "故..." 系列)
            # -------------------------------------------------------
            is_force_line = False
            
            # 4.1 强力包含检查 (新增)
            for kw in STRONG_DELETE_CONTAIN:
                if kw in text:
                    is_force_line = True
                    break
            
            # 4.2 前缀检查 (原有)
            if not is_force_line:
                for prefix in FORCE_DELETE_PREFIXES:
                    if text.startswith(prefix):
                        is_force_line = True
                        break
            
            if is_force_line:
                p.text = ""
                # 注意：这里只删当前行，不改变 is_deleting 状态
                continue

            # -------------------------------------------------------
            # 5. 执行常规删除
            # -------------------------------------------------------
            if is_deleting:
                p.text = "" # 彻底清空内容，包括图片 run

    # === 第二轮：物理清理 (删除空段落) ===
    # 这一步能把逻辑删除留下的空行，以及原本的空行都删掉
    # 同时保护剩下的图片 (必须是 Question 或 Material 里的图片)
    
    count_deleted = 0
    to_delete = []

    # 只处理 Body 这一层的 Paragraph
    for p in doc.paragraphs:
        if not p.text.strip():
            # 只有当既没有文字，也没有图片时，才删除
            # 经过第一轮，Answer 里的图片所在的段落，如果 p.text="" 执行成功，图片应该没了？
            # 验证：p.text="" 会清除段落中所有 run，包括图片。
            # 所以第一轮 is_deleting 为 True 时，p.text="" 已经杀死了图片。
            # 这里剩下的图片应该是 Question / Material 里的。
            if not has_image(p):
                to_delete.append(p)
    
    for p in to_delete:
        p_element = p._element
        if p_element.getparent() is not None:
            p_element.getparent().remove(p_element)
            count_deleted += 1

    doc.save(output_path)
    print(f"  -> 完成。清理空段落数: {count_deleted}")

# ================= 批量入口 =================
def batch_run(folder_path):
    # 查找所有带“解析”的文件
    files = [f for f in os.listdir(folder_path) if '解析' in f and f.endswith('.docx')]
    if not files:
        print("当前目录下未找到带'解析'的docx文件。")
        return

    print(f"找到 {len(files)} 个文件，开始处理...")
    for f in files:
        in_path = os.path.join(folder_path, f)
        # 生成输出文件名：去除“解析”
        out_name = f.replace("-解析", "").replace("解析", "")
        # 如果名字没变（原名里可能没解析或者位置不对），加个前缀防覆盖
        if out_name == f: 
            out_name = "题目版_" + f
            
        out_path = os.path.join(folder_path, out_name)
        
        try:
            clean_docx_complete(in_path, out_path)
        except Exception as e:
            print(f"处理出错 {f}: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    # 默认运行当前目录
    batch_run('.')
