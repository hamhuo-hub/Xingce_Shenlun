import os
import re
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

# ==================== 配置区域 ====================

# 1. 触发【开始删除】的关键词
START_KEYWORDS = ['【答案】', '【解析】', '【拓展】', '【来源】', '正确答案', '参考答案']

# 2. 【强制删除】的前缀特征
# 即使脚本判断“已经进入下一题了”，如果遇到这些开头，依然强制删掉
# 这能解决“误判停止”后，残留的尾部解析
FORCE_DELETE_PREFIXES = [
    '因此，选择', '因此选择', '故本题选', '故正确答案',
    '第一步，', '第二步，', '第三步，', '第四步，',
    'A项：', 'B项：', 'C项：', 'D项：', 'A项 ', 'B项 ',
    '①', '②', '③', '④' # 解析里常出现的序号
]

# 3. 题目序号正则
Q_PATTERN = re.compile(r'^\s*(\d+)[、\.．\)\s]')

# 4. 大标题/材料保护正则 (遇到这些绝对停止删除)
HEADER_PATTERN = re.compile(
    r'^\s*第[一二三四五六七八九十]+部分|'  # 第一部分
    r'^\s*[一二三四五六七八九十]+、|'      # 一、常识判断
    r'^\s*根据.*(材料|回答|短文)'          # 根据...回答
)

# ================================================

def iter_block_items(parent):
    """递归遍历文档内容"""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Parent object error")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def check_start_trigger(text):
    """检查是否触发删除模式"""
    for kw in START_KEYWORDS:
        if kw in text:
            return kw
    return None

def is_valid_next_question(found_num, last_num):
    """
    【核心逻辑】验证找到的数字是不是合理的下一题序号
    防止解析里的 "1. xxx" 打断删除
    """
    # 1. 如果还没找到过题目 (last_num=0)，那 1 肯定是题目
    if last_num == 0:
        return True
    
    # 2. 如果找到的数字 比 上一题数字 小很多 (例如上一题100，这里发现了1)，大概率是解析里的列表
    if found_num < last_num:
        # 特例：除非它重置为1了（可能是新的一部分），但通常新部分会有大标题触发 HEADER_PATTERN
        # 这里保守一点：如果是1，且距离上一题很远，认为是干扰
        if found_num == 1 and last_num > 10:
            return False
        return False # 乱序或回退，视为干扰
        
    # 3. 如果跳跃太大 (例如上一题5，这里发现了100)，可能是误判，但在行测里不常见
    # 允许适度的跳号（漏题），但不允许太离谱的
    if found_num - last_num > 20:
        return False
        
    return True

def clean_docx_smart(input_path, output_path):
    print(f"处理文件: {os.path.basename(input_path)}")
    doc = Document(input_path)
    
    # 状态变量
    is_deleting = False
    last_q_num = 0  # 记录上一个确认为题目的序号
    
    # 获取所有块
    blocks = list(iter_block_items(doc))
    
    # === 第一轮：智能逻辑擦除 ===
    for block in blocks:
        # 统一获取需要处理的段落对象列表
        paragraphs = []
        if isinstance(block, Paragraph):
            paragraphs.append(block)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        
        for p in paragraphs:
            text = p.text.strip()
            if not text: continue

            # -------------------------------------------------------
            # 1. 优先检查：大标题/材料 (最高优先级停止删除)
            # -------------------------------------------------------
            if HEADER_PATTERN.match(text):
                if is_deleting:
                    # print(f"  [保护] 遇到大标题，停止删除: {text[:10]}...")
                    pass
                is_deleting = False
                last_q_num = 0 # 重置计数器（可选，视试卷结构而定）
                continue

            # -------------------------------------------------------
            # 2. 检查：看似是题目序号
            # -------------------------------------------------------
            match = Q_PATTERN.match(text)
            if match:
                try:
                    found_num = int(match.group(1))
                    
                    # 【关键改进】：逻辑验证
                    # 只有当这个数字 "合理" 时，才停止删除
                    if is_valid_next_question(found_num, last_q_num):
                        is_deleting = False
                        last_q_num = found_num
                        # print(f"  [题目] 确认第 {found_num} 题")
                    else:
                        # 虽然像题目(如 "1.")，但不合理，视为解析的一部分，保持删除状态
                        # print(f"  [过滤] 忽略解析中的干扰序号: {found_num} (上一题: {last_q_num})")
                        pass
                except ValueError:
                    pass

            # -------------------------------------------------------
            # 3. 检查：开始删除触发器 (答案/解析)
            # -------------------------------------------------------
            keyword = check_start_trigger(text)
            if keyword:
                is_deleting = True
                if text.startswith(keyword):
                    p.text = ""
                else:
                    parts = p.text.split(keyword)
                    p.text = parts[0].strip()
                continue

            # -------------------------------------------------------
            # 4. 检查：强制删除特征 (补刀)
            # -------------------------------------------------------
            # 即使 is_deleting == False，如果这行长得像解析，也删
            is_force_line = False
            for prefix in FORCE_DELETE_PREFIXES:
                if text.startswith(prefix):
                    is_force_line = True
                    break
            
            if is_force_line:
                # print(f"  [补刀] 强制删除残留解析: {text[:10]}...")
                p.text = ""
                continue

            # -------------------------------------------------------
            # 5. 执行常规删除
            # -------------------------------------------------------
            if is_deleting:
                p.text = ""

    # === 第二轮：物理清理 (删除空段落) ===
    count_deleted = 0
    
    # 重新遍历所有段落
    all_paragraphs = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            all_paragraphs.append(block)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

    for p in all_paragraphs:
        if not p.text.strip():
            # 图片保护
            has_pic = p._element.findall('.//w:drawing', namespaces=p._element.nsmap) or \
                      p._element.findall('.//w:pict', namespaces=p._element.nsmap)
            
            if not has_pic:
                p_element = p._element
                if p_element.getparent() is not None:
                    p_element.getparent().remove(p_element)
                    count_deleted += 1

    doc.save(output_path)
    print(f"  -> 完成。清理段落数: {count_deleted}")

# ================= 批量入口 =================
def batch_run(folder_path):
    files = [f for f in os.listdir(folder_path) if '解析' in f and f.endswith('.docx')]
    if not files:
        print("未找到文件。")
        return

    for f in files:
        in_path = os.path.join(folder_path, f)
        out_name = f.replace("-解析", "").replace("解析", "")
        if out_name == f: out_name = "题目版_" + f
        out_path = os.path.join(folder_path, out_name)
        
        try:
            clean_docx_smart(in_path, out_path)
        except Exception as e:
            print(f"错误 {f}: {e}")

if __name__ == "__main__":
    batch_run('.')