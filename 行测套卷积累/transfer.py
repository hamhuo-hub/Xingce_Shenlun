import re
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def has_image(paragraph):
    """
    检查段落是否包含图片 (drawing 或 pict)
    如果有图片，绝对不能删
    """
    drawings = paragraph._element.findall('.//w:drawing', namespaces=paragraph._element.nsmap)
    if drawings: return True
    picts = paragraph._element.findall('.//w:pict', namespaces=paragraph._element.nsmap)
    if picts: return True
    return False

def iter_block_items(parent):
    """
    生成器：按顺序获取文档中的段落和表格
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def clean_docx_final_v2(input_path, output_path):
    doc = Document(input_path)
    
    # ==========================================
    # 0. 特殊处理：第一行标题去除“解析”二字
    # ==========================================
    if len(doc.paragraphs) > 0:
        title_p = doc.paragraphs[0]
        if "解析" in title_p.text:
            print(f"处理标题: '{title_p.text.strip()}' -> 去除'解析'")
            title_p.text = title_p.text.replace("解析", "")

    next_q_num = 1
    delete_mode = False
    
    # 正则：题目序号 (数字+顿号/点)
    q_pattern = re.compile(r'^\s*(\d+)\s*[、\.．]')
    
    # 正则：材料/标题保护 (遇到这些特征强制停止删除)
    # 新增：|^\s*第[一二三四五六七八九十]+部分 -> 匹配 "第一部分 常识判断"
    material_pattern = re.compile(r'^\s*[一二三四五六七八九十]+、|^\s*第[一二三四五六七八九十]+部分|.*根据.*材料|.*回答.*[0-9]+.*题|.*阅读.*短文')

    print(f"开始深度清理: {input_path}")
    
    blocks = list(iter_block_items(doc))
    
    # ==========================================
    # 第一轮：逻辑擦除 (只清空内容，不删对象)
    # ==========================================
    for block in blocks:
        if isinstance(block, Paragraph):
            text = block.text.strip()
            
            # 1. 检查【题目序号】 -> 停止删除
            match = q_pattern.match(text)
            if match:
                try:
                    num_found = int(match.group(1))
                    if num_found == next_q_num:
                        delete_mode = False
                        next_q_num += 1
                        continue 
                except ValueError:
                    pass

            # 2. 检查【材料/大标题】 -> 停止删除 (含第n部分)
            # 即使处于删除模式，遇到大标题也要立即醒来
            if material_pattern.match(text):
                if delete_mode:
                    print(f"保留大标题/材料: {text[:15]}...")
                delete_mode = False
                continue

            # 3. 检查【答案】 -> 开启删除
            if '【答案】' in text:
                delete_mode = True
                if text.startswith('【答案】'):
                    block.text = "" # 清空整段
                else:
                    parts = block.text.split('【答案】')
                    block.text = parts[0].strip()
                continue

            # 4. 执行逻辑删除
            if delete_mode:
                block.text = "" 

        elif isinstance(block, Table):
            table_text = ""
            for row in block.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
            table_text = table_text.strip()
            
            match = q_pattern.match(table_text)
            if match:
                try:
                    num_found = int(match.group(1))
                    if num_found == next_q_num:
                        delete_mode = False
                        next_q_num += 1
                        continue 
                except ValueError:
                    pass

            if '【答案】' in table_text:
                delete_mode = True
                pass 

            if delete_mode:
                tbl_element = block._element
                if tbl_element.getparent() is not None:
                    tbl_element.getparent().remove(tbl_element)

    # ==========================================
    # 第二轮：物理清洗 (寸草不生版)
    # ==========================================
    paragraphs_to_remove = []
    
    for p in doc.paragraphs:
        txt = p.text.strip()
        # 只要没字且没图，直接删 (无视分节符、分页符)
        if not txt: 
            if not has_image(p):
                paragraphs_to_remove.append(p)
    
    print(f"扫描到 {len(paragraphs_to_remove)} 个空段落，正在移除...")
    
    for p in paragraphs_to_remove:
        p_element = p._element
        if p_element.getparent() is not None:
            p_element.getparent().remove(p_element)

    doc.save(output_path)
    print(f"清理完毕！文件已保存至: {output_path}")

if __name__ == "__main__":
    input_file = "行测组卷8-解析.docx"
    output_file = "行测组卷8.docx"
    
    try:
        clean_docx_final_v2(input_file, output_file)
    except Exception as e:
        print(f"运行出错: {e}")