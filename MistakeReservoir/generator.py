import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup

class PaperBuilder:
    def __init__(self, media_dir: str):
        self.media_dir = media_dir
        
    def create_paper(self, questions: list, output_path: str):
        doc = Document()
        
        # --- Styles ---
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(10.5) # 5号
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0
        
        # --- Logic ---
        # Sort Order
        type_map = {
            '常识': 1, 
            '言语': 2, 
            '数量': 3, 
            '判断': 4, '图形': 4.1, '定义': 4.2, '类比': 4.3, '逻辑': 4.4,
            '资料': 5
        }
        
        questions.sort(key=lambda q: (type_map.get(q.get('type', ''), 99), q.get('material_id') or 0, q.get('original_num')))
        
        current_type = None
        last_material_id = None
        global_idx = 1
        
        section_map = {
            '常识': '第一部分 常识判断',
            '言语': '第二部分 言语理解与表达',
            '数量': '第三部分 数量关系',
            '判断': '第四部分 判断推理',
            '图形': '第四部分 判断推理', # Subtypes grouped under main
            '定义': '第四部分 判断推理',
            '类比': '第四部分 判断推理',
            '逻辑': '第四部分 判断推理',
            '资料': '第五部分 资料分析'
        }
        
        # Track main sections added to avoid repeats
        added_encounters = set()
        
        for q in questions:
            q_type = q.get('type') or "其他"
            clean_type = q_type
            
            # Map subtype to main type for header
            if q_type in ['图形', '定义', '类比', '逻辑']:
                clean_type = '判断'
                
            header_title = section_map.get(clean_type, f"部分 {clean_type}")
            
            # 1. Section Header
            if header_title not in added_encounters:
                if added_encounters: doc.add_paragraph() # Spacing
                h = doc.add_heading(header_title, level=1)
                h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                added_encounters.add(header_title)
                current_type = clean_type # Logical group
                last_material_id = None
            
            # 2. Material
            mid = q.get('material_id')
            if mid and mid != last_material_id:
                mat_content = q.get('material_content')
                if mat_content:
                    p = doc.add_paragraph()
                    r = p.add_run("根据以下材料，回答下列问题：")
                    r.bold = True
                    self._add_html_content(doc, mat_content)
                    doc.add_paragraph() # Spacing after material
                last_material_id = mid
            elif not mid:
                last_material_id = None

            # 3. Question Logic
            # Stem
            p = doc.add_paragraph()
            p.add_run(f"{global_idx}. ").bold = True
            self._add_html_content_inline(p, q.get('content_html'), doc)
            
            # Options
            if q.get('options_html'):
                 self._add_options(doc, q.get('options_html'))
            
            global_idx += 1
            doc.add_paragraph() # Spacing between questions

        # --- Answer Key ---
        doc.add_page_break()
        h = doc.add_heading('参考答案与解析', level=1)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for i, q in enumerate(questions):
            p = doc.add_paragraph()
            p.add_run(f"{i+1}. ").bold = True
            
            ans = q.get('answer_html')
            if ans:
                self._add_html_content_inline(p, ans, doc)
            else:
                p.add_run("（暂无解析）")
                
        doc.save(output_path)
        return output_path

    def _add_html_content(self, doc, html_str):
        """Block level adder"""
        if not html_str: return
        soup = BeautifulSoup(html_str, 'html.parser')
        
        # Extract images and text
        # If div.img-container -> Image
        # If p -> text
        
        for elem in soup.find_all(['p', 'div', 'table']):
            if elem.name == 'p':
                if elem.get_text().strip():
                    doc.add_paragraph(elem.get_text().strip())
            elif elem.name == 'div' and 'img-container' in elem.get('class', []):
                self._add_image(doc, elem)
            elif elem.name == 'table':
                # Simplified table
                pass

    def _add_html_content_inline(self, paragraph, html_str, doc):
        """Adds text to existing paragraph, inserting images inline or as blocks based on size"""
        if not html_str: return
        soup = BeautifulSoup(html_str, 'html.parser')
        
        # Iterate over child nodes to maintain order
        # Note: This is a simplistic traversal. Nested tags might need recursion, 
        # but usually the input HTML is flat-ish (p, img, span).
        
        # Determine the run to append to
        run = paragraph.add_run()
        
        # We process 'descendants' carefully or just iterate contents?
        # contents is strictly direct children. text might be split.
        
        # Strategy: Flatten the soup to a list of (type, content)
        # Text -> append to current run
        # Img -> check size -> append to run OR break paragraph
        
        # Simple recursive walker (flattened)
        for output in self._flatten_nodes(soup):
            type_, content = output
            if type_ == 'text':
                if content: run.add_text(content)
            elif type_ == 'img':
                # Try to insert
                self._insert_image_hybrid(doc, run, content)
                
    def _flatten_nodes(self, element):
        """Yields ('text', str) or ('img', src)"""
        if element.name == 'img':
            yield ('img', element.get('src'))
            return

        if isinstance(element, str): # NavigableString
            yield ('text', str(element))
            return

        for child in element.children:
            yield from self._flatten_nodes(child)

    def _add_image(self, doc, elem):
        img = elem.find('img')
        if img:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            self._insert_image_hybrid(doc, run, img.get('src'))

    def _add_options(self, doc, html_str):
        """Parses options and puts them on new lines, preserving images"""
        if not html_str: return
        soup = BeautifulSoup(html_str, 'html.parser')
        
        # If wrapped in <p>, iterate Ps. 
        # Else, just parse the whole blob.
        ps = soup.find_all('p')
        
        if ps:
            for p_tag in ps:
                p = doc.add_paragraph()
                self._add_html_content_inline(p, str(p_tag), doc)
        else:
            # Fallback for unwrapped text
            p = doc.add_paragraph()
            self._add_html_content_inline(p, html_str, doc)

    def _insert_image_hybrid(self, doc, run, src):
        """
        Inserts image. 
        - If small/icon-like: Insert into 'run' with height=Pt(11) (Inline).
        - If large: Insert as new Paragraph (Block).
        """
        try:
            from PIL import Image
        except ImportError:
            Image = None

        if not src: return
        fname = src.split('/')[-1]
        fpath = os.path.join(self.media_dir, fname)
        
        if not os.path.exists(fpath):
            print(f"DEBUG: Image missing {fpath}")
            return
            
        # Determine Sizing Strategy
        is_inline = False
        width_arg = None
        height_arg = None
        
        if Image:
            try:
                with Image.open(fpath) as img:
                    w, h = img.size
                    
                    # Revised Heuristic for "Small / Inline"
                    # User specifically wants images to match 5-hao font (~10.5pt, approx 14-20px rendered).
                    # If an image is "relatively small" (e.g. < 250px height), assume it's an inline symbol/formula and shrink it.
                    # 250px is arbitrary but covers most high-dpi small icons.
                    
                    if h < 250: 
                        is_inline = True
                        height_arg = Pt(11) # Force to 5-hao size
                    else:
                        # Large content (Chart, Screenshot)
                        is_inline = False
                        if w > 400:
                            width_arg = Inches(5.5) # Max Page Width
                        else:
                             width_arg = Inches(3.5) if w > 300 else None
            except:
                is_inline = False # Fallback to block on error
                width_arg = Inches(2.0)
        
        if is_inline:
            # Add to CURRENT run
            try:
                run.add_picture(fpath, height=height_arg)
            except Exception as e:
                 print(f"Error adding inline pic: {e}")
        else:
            # Add to NEW paragraph
            # We need to break the flow? 
            # Ideally we'd close the current run, make a new p, then resume?
            # But we are inside `_add_html_content_inline` taking a `paragraph`.
            # We can't easily "split" the paragraph passed in unless we return new context.
            # Workaround: Add to the *End* of the current paragraph via run?
            # run.add_picture() adds it at the current position. 
            # If we want a "Block" feel but are stuck in a run, we can add a break before/after?
            try:
                run.add_break()
                if width_arg:
                    run.add_picture(fpath, width=width_arg)
                else:
                    run.add_picture(fpath)
                run.add_break()
            except Exception as e:
                pass
