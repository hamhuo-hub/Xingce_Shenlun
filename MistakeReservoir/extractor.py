import os
import re
import uuid
import shutil
from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

class QuestionExtractor:
    def __init__(self, media_dir: str):
        self.media_dir = media_dir
        if not os.path.exists(media_dir):
            os.makedirs(media_dir)
            
        # Regex Patterns
        self.Q_PATTERN = re.compile(r'^\s*\(?\d+\)?[\.．、\s]')
        self.HEADER_PATTERN = re.compile(
            r'^\s*第[一二三四五六七八九十]+部分|'
            r'^\s*[一二三四五六七八九十]+、|'
            r'^\s*(根据|阅读).*(材料|回答|短文)'
        )
        # Unified Answer Regex (covers spaces and various formats)
        self.ANSWER_REGEX = re.compile(
            r'(【\s*答案\s*】|【\s*解析\s*】|【\s*拓展\s*】|【\s*来源\s*】|正确\s*答案|参考\s*答案|答案\s*[:：]|解析\s*[:：])'
        )
        self.OPTION_PATTERN = re.compile(r'^\s*\(?[A-D]\)?[\.．、\s]')
        
        # Current State
        self.current_material_id = None
        self.current_material_content = ""
        self.current_type = "Unknown"

    def iter_block_items(self, parent):
        """Iterate through docx blocks (Paragraphs and Tables)"""
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("Parent object error")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def extract_images(self, block) -> List[str]:
        """
        Extract images from a block (Paragraph/Table)
        Save to media_dir and return list of filenames.
        """
        image_paths = []
        
        # Access the underlying xml element
        if isinstance(block, Paragraph):
            elms = [block._element]
        elif isinstance(block, Table):
            elms = []
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        elms.append(p._element)
        
        for elm in elms:
            pass 
        return image_paths

    def _save_image_from_blip(self, doc, blip_rId) -> Optional[str]:
        """
        Internal: Save image binary from blip relationship ID
        """
        try:
            if not blip_rId: return None
            
            # Safe access to related_parts
            if blip_rId not in doc.part.related_parts:
                return None
                
            image_part = doc.part.related_parts[blip_rId]
            # Generate filename
            try:
                ext = image_part.content_type.split('/')[-1]
                if ext == 'jpeg': ext = 'jpg'
            except:
                ext = 'png'
            
            filename = f"{uuid.uuid4().hex}.{ext}"
            filepath = os.path.join(self.media_dir, filename)
            
            with open(filepath, "wb") as f:
                f.write(image_part.blob)
            
            return filename
        except Exception as e:
            print(f"Error saving image {blip_rId}: {e}")
            return None

    def get_block_images(self, doc, block) -> List[str]:
        """Real implementation of image extraction for a block"""
        images = []
        try:
            if isinstance(block, Paragraph):
                # Search for <a:blip>
                ns = block._element.nsmap
                # Ensure 'a' prefix exists if we use it, or use the full namespace
                if 'a' not in ns:
                    ns['a'] = 'http://schemas.openxmlformats.org/drawingml/2006/main'
                
                try:
                    blips = block._element.findall('.//a:blip', ns)
                except KeyError:
                    blips = []

                for blip in blips:
                    rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rId:
                        fname = self._save_image_from_blip(doc, rId)
                        if fname: images.append(fname)
                        
                # check v:imagedata
                imagedatas = block._element.findall('.//v:imagedata', ns)
                for idata in imagedatas:
                    rId = idata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if rId:
                        fname = self._save_image_from_blip(doc, rId)
                        if fname: images.append(fname)
                        
            elif isinstance(block, Table):
                for row in block.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            images.extend(self.get_block_images(doc, p))
        except Exception as e:
            pass
            
        return images

    def block_to_html(self, doc, block, skip_images=False) -> Tuple[str, List[str]]:
        """Convert block to simple HTML and extract images"""
        images = [] if skip_images else self.get_block_images(doc, block)
        html = ""
        
        if isinstance(block, Paragraph):
            text = block.text.strip()
            html = f"<p>{text}</p>" if text else ""
        elif isinstance(block, Table):
            rows = []
            for row in block.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cells.append(f"<td>{cell_text}</td>")
                rows.append(f"<tr>{''.join(cells)}</tr>")
            html = f"<table border='1' cellspacing='0' cellpadding='5'>{''.join(rows)}</table>"
        
        for img in images:
            html += f'<div class="img-container"><img src="/media/{img}" class="question-img" /></div>'
                
        return html, images

    def process_buffer_as_question(self, doc, buffer: List, q_num: int, skip_images=False) -> Dict:
        """
        Convert a buffer of blocks into structured Question data.
        Separates Stem, Options, and Analysis.
        """
        from copy import deepcopy
        
        stem_blocks = []
        option_blocks = []
        analysis_blocks = []
        
        # State: 0=Stem, 1=Options, 2=Analysis
        state = 0
        
        for block in buffer:
            text = ""
            if isinstance(block, Paragraph):
                text = block.text.strip()
            elif isinstance(block, Table):
                # Extract text from table for keyword checking
                cell_texts = []
                for row in block.rows:
                    for cell in row.cells:
                        cell_texts.append(cell.text.strip())
                text = " ".join(cell_texts)
            
            # Check Switch to Analysis using Regex
            # Re-defined pattern string locally to ensure no corruption
            ans_pattern = r'(【\s*答案\s*】|【\s*解析\s*】|【\s*拓展\s*】|【\s*来源\s*】|正确\s*答案|参考\s*答案|答案\s*[:：]|解析\s*[:：])'
            ans_match = re.search(ans_pattern, text)
            
            if ans_match:
                start_idx = ans_match.start()
                
                if start_idx == 0:
                    state = 2
                else:
                    if isinstance(block, Paragraph):
                        part1_text = text[:start_idx].strip()
                        part2_text = text[start_idx:].strip()
                        
                        try:
                            elem_copy = deepcopy(block._element)
                            block_part2 = Paragraph(elem_copy, block._parent)
                            block_part2.text = part2_text
                            
                            block.text = part1_text
                            
                            if state == 1:
                                option_blocks.append(block)
                            elif state == 2:
                                analysis_blocks.append(block)
                            else:
                                stem_blocks.append(block)
                                
                            state = 2
                            analysis_blocks.append(block_part2)
                            
                            continue 
                            
                        except Exception as e:
                            print(f"Wrapper Split Error: {e}")
                            state = 2 
                    else:
                        state = 2 
            
            if state < 2:
                # Direct regex check for options
                # Matches A. B. C. D. at start of line
                opt_pattern = r'^\s*\(?[A-D]\)?[\.．、\s]'
                if re.match(opt_pattern, text):
                    state = 1
            
            if state == 2:
                analysis_blocks.append(block)
            elif state == 1:
                option_blocks.append(block)
            else:
                stem_blocks.append(block)
        
        def blocks_to_html_str(blks, is_stem=False):
            htmls = []
            imgs = []
            for i_idx, b in enumerate(blks):
                if is_stem and i_idx == 0 and isinstance(b, Paragraph):
                    text = b.text.strip()
                    match = self.Q_PATTERN.match(text)
                    if match:
                        cleaned_text = text[match.end():].strip()
                        block_imgs = [] if skip_images else self.get_block_images(doc, b)
                        imgs.extend(block_imgs)
                        
                        h = f"<p>{cleaned_text}</p>" if cleaned_text else ""
                        
                        for img in block_imgs:
                            h += f'<div class="img-container"><img src="/media/{img}" class="question-img" /></div>'
                        
                        htmls.append(h)
                        continue
                
                h, i = self.block_to_html(doc, b, skip_images=skip_images)
                htmls.append(h)
                imgs.extend(i)
            return "".join(htmls), imgs

        stem_html, stem_imgs = blocks_to_html_str(stem_blocks, is_stem=True)
        opt_html, opt_imgs = blocks_to_html_str(option_blocks)
        ana_html, ana_imgs = blocks_to_html_str(analysis_blocks)
        
        return {
            "original_num": q_num,
            "content_html": stem_html,
            "options_html": opt_html,
            "answer_html": ana_html, 
            "images": stem_imgs + opt_imgs + ana_imgs,
            "type": self.current_type,
            "material_content": self.current_material_content if self.current_material_content else None
        }

    def extract_from_file(self, docx_path: str, target_ids: List[int] = None, skip_images: bool = False) -> List[Dict]:
        """
        Main Enty: Parse file and return list of Question Dicts.
        If target_ids is None, return all.
        """
        doc = Document(docx_path)
        blocks = list(self.iter_block_items(doc))
        
        extracted_questions = []
        buffer = []
        last_q_num = 0 # Track previous question number across sections
        current_q_num = 0
        
        FORCE_DELETE_LINES = ['故', '故。', '故本题选', '故正确答案']
        
        for block in blocks:
            text = ""
            if isinstance(block, Paragraph):
                text = block.text.strip()
            elif isinstance(block, Table):
                 pass

            # 1. Check Header (Material / Type Change)
            if self.HEADER_PATTERN.match(text):
                if buffer and current_q_num > 0:
                    q = self.process_buffer_as_question(doc, buffer, current_q_num, skip_images=skip_images)
                    if target_ids is None or current_q_num in target_ids:
                         extracted_questions.append(q)
                    
                    buffer = []
                
                # Identify Type
                is_material_header = text.strip().startswith("根据") or text.strip().startswith("阅读")
                if not is_material_header:
                    if "常识" in text: self.current_type = "常识"
                    elif "言语" in text: self.current_type = "言语"
                    elif "数量" in text: self.current_type = "数量"
                    elif "资料" in text: self.current_type = "资料"
                    elif "判断" in text: self.current_type = "判断" 
                    
                    if "图形" in text and "推理" in text: self.current_type = "图形"
                    elif "定义" in text and "判断" in text: self.current_type = "定义"
                    elif "类比" in text and "推理" in text: self.current_type = "类比"
                    elif "逻辑" in text and "判断" in text: self.current_type = "逻辑"
                
                # Material Handling
                # Save state and Reset current to 0 to enter "Material Mode"
                if current_q_num > 0:
                    last_q_num = current_q_num
                
                current_q_num = 0
                self.current_material_content = "" 
                
                if "根据" in text or "材料" in text or "阅读" in text:
                     h, _ = self.block_to_html(doc, block, skip_images=skip_images)
                     self.current_material_content += h
                
                continue

            # 2. Check Question Start
            match = self.Q_PATTERN.match(text)
            is_new_q = False
            found_num = 0
            
            if match:
                try:
                    nums = re.findall(r'\d+', text)
                    if nums:
                        found_num = int(nums[0])
                        # Sanity Check
                        if found_num < 500: # Years like 2016 filtered
                            if current_q_num == 0:
                                # Check against LAST q_num if current is 0
                                if last_q_num == 0:
                                    is_new_q = True # Start of file
                                elif (found_num == last_q_num + 1) or (found_num > last_q_num and found_num - last_q_num < 20):
                                     is_new_q = True
                            elif (found_num == current_q_num + 1) or (found_num > current_q_num and found_num - current_q_num < 20):
                                is_new_q = True
                except:
                    pass
            
            if is_new_q:
                # Process previous
                if buffer:
                    if current_q_num > 0:
                        q = self.process_buffer_as_question(doc, buffer, current_q_num, skip_images=skip_images)
                        if target_ids is None or current_q_num in target_ids:
                            extracted_questions.append(q)
                    else:
                        for b in buffer:
                            h, imgs = self.block_to_html(doc, b, skip_images=skip_images)
                            self.current_material_content += h
                
                # Start new
                current_q_num = found_num
                buffer = [block]
                
            else:
                if current_q_num > 0:
                    should_skip = False
                    if isinstance(block, Paragraph) and text in FORCE_DELETE_LINES:
                        should_skip = True
                    
                    if not should_skip:
                        buffer.append(block)
                else:
                    h, imgs = self.block_to_html(doc, block, skip_images=skip_images)
                    if text or imgs:
                        self.current_material_content += h

        if buffer and current_q_num > 0:
            q = self.process_buffer_as_question(doc, buffer, current_q_num, skip_images=skip_images)
            if target_ids is None or current_q_num in target_ids:
                extracted_questions.append(q)
                
        return extracted_questions

if __name__ == "__main__":
    extractor = QuestionExtractor(media_dir="media")
